import pandas as pd
import re
from collections import Counter
from docx import Document
from docx.shared import Pt

# === File Paths ===
search_console_path = r"C:\Users\ksaur\Downloads\vypzee.com-Performance-on-Search-2025-06-06.xlsx"
organic_positions_path = r"C:\Users\ksaur\Downloads\vypzee.com-organic.Positions-in-20250605-2025-06-06T05_09_51Z.xlsx"
output_docx = r"C:\Users\ksaur\Downloads\vypzee_blog_plan_final.docx"

# === Load Data ===
sc_data = pd.ExcelFile(search_console_path).parse("Queries").dropna(subset=["Top queries"])
organic_data = pd.read_excel(organic_positions_path)

# === Data Cleanup ===
sc_data["CTR"] = sc_data["CTR"].astype(float)
sc_data["Impressions"] = sc_data["Impressions"].astype(int)
sc_data["Clicks"] = sc_data["Clicks"].astype(int)
sc_data["Position"] = sc_data["Position"].astype(float)
sc_data["Top queries"] = sc_data["Top queries"].astype(str)

organic_data.columns = organic_data.columns.str.strip()
organic_data = organic_data.dropna(subset=["Keyword"])
organic_data["Keyword"] = organic_data["Keyword"].astype(str)

# === Start Document ===
doc = Document()
doc.add_heading('Vypzee SEO Blog Strategy & Keyword Insights – June 2025', 0)

def add_section(title, items):
    doc.add_heading(title, level=1)
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

# === Section 1: High-Impression, Low-CTR ===Improve meta titles and meta descriptions
low_ctr = sc_data[(sc_data["Impressions"] > 500) & (sc_data["CTR"] < 0.02)]
section = [f"'{row['Top queries']}' (Impr: {row['Impressions']}, CTR: {row['CTR']*100:.2f}%)"
           for _, row in low_ctr.sort_values("Impressions", ascending=False).head(20).iterrows()]
add_section("1. High-Impression but Low-CTR Queries", section)

# === Section 2: Buyer Intent Keywords ===
buyer = sc_data[sc_data["Top queries"].str.contains("buy|price|near me|best|online|cost", case=False)]
section = [f"'{row['Top queries']}' (Clicks: {row['Clicks']}, Pos: {row['Position']:.1f})"
           for _, row in buyer.sort_values("Clicks", ascending=False).head(20).iterrows()]
add_section("2. Buyer-Intent Keywords for Transactional Blog Posts", section)

# === Section 3: Long-Tail Keywords ===Add these keywords as secondary keywords in existing pages.
queries = sc_data["Top queries"].tolist()
long_tail = [q for q in queries if len(q.split()) >= 4][:20]
add_section("3. Long-Tail Blog Opportunities", long_tail)

# === Section 4: Suggested Blog Titles ===
titles = []
for q in long_tail:
    if "near me" in q.lower():
        titles.append(f"Where to Buy {q.title().replace('Near Me', '')} Near You?")
    elif "buy" in q.lower():
        titles.append(f"How to Buy {q.title()} at the Best Price")
    else:
        titles.append(f"Ultimate Guide to {q.title()} in 2025")
add_section("4. Suggested Blog Titles", titles)

# === Section 5: Content Gaps ===improve on-page SEO
content_gap = sc_data[(sc_data["Impressions"] > 300) & (sc_data["Position"] > 10)]
section = [f"'{row['Top queries']}' (Impr: {row['Impressions']}, Pos: {row['Position']:.1f})"
           for _, row in content_gap.sort_values("Impressions", ascending=False).head(15).iterrows()]
add_section("6. Content Gaps: High Impressions, Poor Rankings", section)

# === Section 6: Low-Hanging Fruits ===Focus link-building
almost = sc_data[(sc_data["Position"] >= 5) & (sc_data["Position"] <= 15)]
section = [f"'{row['Top queries']}' (Pos: {row['Position']:.1f}, CTR: {row['CTR']*100:.2f}%)"
           for _, row in almost.sort_values("Position").head(10).iterrows()]
add_section("8. Low-Hanging Fruits: Page 1 but Not Top", section)

# === Section 8: Geo Keywords ===Create dedicated local market pages
geo_kw = sc_data[sc_data["Top queries"].str.contains("delhi|noida|gurgaon", case=False)]
section = [f"'{row['Top queries']}' (Clicks: {row['Clicks']}, Impr: {row['Impressions']})"
           for _, row in geo_kw.head(15).iterrows()]
add_section("11. Geo-Based Local Search Keywords", section)

# === Section 9: Top Organic Keywords (from Positions file) ===
top_organic = organic_data[["Keyword", "Position", "Search Volume", "Traffic", "Traffic (%)"]].dropna().sort_values("Traffic", ascending=False).head(15)
section = [f"{row['Keyword']} - Pos: {row['Position']}, Volume: {row['Search Volume']}, Traffic: {row['Traffic']} ({row['Traffic (%)']}%)"
           for _, row in top_organic.iterrows()]
add_section("13. Top Organic Traffic Keywords", section)
print("strong internal linking &fresh with updated content")

# === Section 10: Underperforming CTRs vs Expected ===
def expected_ctr(pos):
    if pos <= 1: return 0.3
    elif pos <= 2: return 0.15
    elif pos <= 3: return 0.1
    elif pos <= 6: return 0.05
    else: return 0.02

sc_data["Expected CTR"] = sc_data["Position"].apply(expected_ctr)
sc_data["CTR Gap"] = sc_data["Expected CTR"] - sc_data["CTR"]
underperforming = sc_data[sc_data["CTR Gap"] > 0.05]
section = [f"{row['Top queries']} (Actual CTR: {row['CTR']:.2%}, Expected: {row['Expected CTR']:.2%})"
           for _, row in underperforming.head(15).iterrows()]
add_section("10. Underperforming CTRs vs Expected", section)

# === Section 11: Clustering Queries by Delhi Markets ===
markets = [
    "chandni chowk", "chawri bazar", "karol bagh", "lajpat nagar",
    "sarojini nagar", "sadar bazar", "kamla nagar", "janpath",
    "palika bazar", "connaught place", "gk market", "dilli haat",
    "nehru place", "shankar market", "south ex", "malviya nagar",
    "gandhi nagar", "bhogal market", "meena bazar", "tilak nagar",
    "khari baoli", "vikas marg", "rajouri garden", "pitampura",
    "model town market", "krishna nagar", "moti nagar", "laxmi nagar",
    "geeta colony market", "green park market", "hauz khas market",
    "defence colony market", "daryaganj market", "mayur vihar market",
    "okhla market", "jamia nagar", "trilokpuri", "anand vihar market",
    "noida sector 18", "noida sector 62", "atta market", "galleria market",
    "gurgaon sector 14 market", "gurgaon huda market", "mg road gurgaon",
    "bhangel market", "indrapuram market", "vaishali market", "kaushambi market",
    "vasundhara market", "rohini sector 7", "rohini sector 9", "rohini west",
    "dwarka sector 6", "dwarka sector 10", "dwarka mor", "janakpuri district center",
    "tilak nagar", "paharganj market", "nangloi market", "bawana market",
    "badarpur market", "jhandewalan market", "seelampur market"
]

market_clusters = {market.title(): [] for market in markets}

# Classify queries into market-based buckets
for query in sc_data["Top queries"]:
    for market in markets:
        if market in query.lower():
            market_clusters[market.title()].append(query)
            break  # avoid double-counting

# Filter out empty clusters and write to doc
doc.add_heading("11. Clustering Queries by Delhi Markets", level=1)
for market, keywords in market_clusters.items():
    if keywords:
        doc.add_heading(market, level=2)
        for kw in keywords[:10]:  # show only top 10 for each
            doc.add_paragraph(kw, style="List Bullet")


# === Save DOCX ===
doc.save(output_docx)
print(f"✅ SEO blog plan saved to: {output_docx}")
